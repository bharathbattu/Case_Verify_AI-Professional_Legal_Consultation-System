"""
Document Processing Module for Case-Verify AI
Handles PDF upload, OCR integration, and template generation
"""

import streamlit as st
import io
import base64
from typing import Optional, Dict, Any, List, Tuple
import logging
from datetime import datetime
import json
import re

# Document processing imports
try:
    import PyPDF2
    import fitz  # PyMuPDF for better PDF handling
    from PIL import Image
    import pytesseract
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    import mammoth
except ImportError as e:
    logging.warning(f"Document processing dependencies not installed: {e}")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Main document processing class"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg']
        self.max_file_size = 10 * 1024 * 1024  # 10MB limit
    
    def process_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """
        Process uploaded file and extract text content
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            if uploaded_file is None:
                return {"success": False, "error": "No file uploaded"}
            
            # Validate file size
            if uploaded_file.size > self.max_file_size:
                return {"success": False, "error": f"File size exceeds {self.max_file_size // (1024*1024)}MB limit"}
            
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if f'.{file_extension}' not in self.supported_formats:
                return {"success": False, "error": f"Unsupported file format: .{file_extension}"}
            
            # Process based on file type
            if file_extension == 'pdf':
                return self._process_pdf(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                return self._process_word_document(uploaded_file)
            elif file_extension == 'txt':
                return self._process_text_file(uploaded_file)
            elif file_extension in ['png', 'jpg', 'jpeg']:
                return self._process_image_ocr(uploaded_file)
            else:
                return {"success": False, "error": "Unsupported file type"}
                
        except Exception as e:
            logger.error(f"Error processing uploaded file: {str(e)}")
            return {"success": False, "error": f"Processing error: {str(e)}"}
    
    def _process_pdf(self, uploaded_file) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            # Read the PDF file
            pdf_bytes = uploaded_file.read()
            
            # Try PyMuPDF first (better for complex PDFs)
            try:
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                text_content = ""
                pages_processed = 0
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text_content += page.get_text()
                    pages_processed += 1
                
                doc.close()
                
                return {
                    "success": True,
                    "text": text_content.strip(),
                    "metadata": {
                        "pages": pages_processed,
                        "file_size": len(pdf_bytes),
                        "processor": "PyMuPDF"
                    }
                }
                
            except Exception:
                # Fallback to PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
                text_content = ""
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text()
                
                return {
                    "success": True,
                    "text": text_content.strip(),
                    "metadata": {
                        "pages": len(pdf_reader.pages),
                        "file_size": len(pdf_bytes),
                        "processor": "PyPDF2"
                    }
                }
                
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            return {"success": False, "error": f"PDF processing failed: {str(e)}"}
    
    def _process_word_document(self, uploaded_file) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            # Read file content
            file_bytes = uploaded_file.read()
            
            if uploaded_file.name.lower().endswith('.docx'):
                # Process DOCX file
                doc = docx.Document(io.BytesIO(file_bytes))
                text_content = ""
                
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                return {
                    "success": True,
                    "text": text_content.strip(),
                    "metadata": {
                        "paragraphs": len(doc.paragraphs),
                        "file_size": len(file_bytes),
                        "processor": "python-docx"
                    }
                }
            else:
                # For .doc files, try mammoth
                try:
                    result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
                    return {
                        "success": True,
                        "text": result.value.strip(),
                        "metadata": {
                            "file_size": len(file_bytes),
                            "processor": "mammoth"
                        }
                    }
                except Exception:
                    return {"success": False, "error": "Unable to process .doc file. Please convert to .docx format."}
                    
        except Exception as e:
            logger.error(f"Word document processing error: {str(e)}")
            return {"success": False, "error": f"Word document processing failed: {str(e)}"}
    
    def _process_text_file(self, uploaded_file) -> Dict[str, Any]:
        """Process plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)
                    text_content = uploaded_file.read().decode(encoding)
                    
                    return {
                        "success": True,
                        "text": text_content.strip(),
                        "metadata": {
                            "encoding": encoding,
                            "file_size": len(text_content.encode()),
                            "processor": "text"
                        }
                    }
                except UnicodeDecodeError:
                    continue
            
            return {"success": False, "error": "Unable to decode text file with supported encodings"}
            
        except Exception as e:
            logger.error(f"Text file processing error: {str(e)}")
            return {"success": False, "error": f"Text file processing failed: {str(e)}"}
    
    def _process_image_ocr(self, uploaded_file) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        try:
            # Read image
            image_bytes = uploaded_file.read()
            image = Image.open(io.BytesIO(image_bytes))
            
            # Perform OCR
            extracted_text = pytesseract.image_to_string(image)
            
            return {
                "success": True,
                "text": extracted_text.strip(),
                "metadata": {
                    "image_size": image.size,
                    "file_size": len(image_bytes),
                    "processor": "tesseract-ocr"
                }
            }
            
        except Exception as e:
            logger.error(f"OCR processing error: {str(e)}")
            return {"success": False, "error": f"OCR processing failed: {str(e)}"}


class LegalTemplateGenerator:
    """Generate legal document templates"""
    
    def __init__(self):
        self.templates = {
            'notice': 'Legal Notice Template',
            'petition': 'Court Petition Template', 
            'application': 'Legal Application Template',
            'affidavit': 'Affidavit Template',
            'complaint': 'Complaint Template'
        }
    
    def generate_legal_notice(self, case_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a legal notice document"""
        try:
            # Create new Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('LEGAL NOTICE', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date and details
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("")
            
            # To section
            doc.add_paragraph("To:")
            doc.add_paragraph("[Name and Address of the Recipient]")
            doc.add_paragraph("")
            
            # Subject
            subject = doc.add_paragraph("Subject: ")
            subject.add_run(f"Legal Notice regarding {case_data.get('relief_type', 'Legal Matter').replace('-', ' ').title()}")
            doc.add_paragraph("")
            
            # Main content
            doc.add_paragraph("Sir/Madam,")
            doc.add_paragraph("")
            
            # Case facts
            facts_para = doc.add_paragraph("WHEREAS, ")
            facts_para.add_run(case_data.get('facts', '[Case facts to be inserted here]'))
            doc.add_paragraph("")
            
            # Legal demand
            doc.add_paragraph("NOW THEREFORE, I hereby call upon you to:")
            
            demands = [
                f"Resolve the matter related to {case_data.get('relief_type', 'the legal issue')} within 15 days",
                "Provide adequate compensation for the damages caused",
                "Cease any further actions that may aggravate the situation"
            ]
            
            for demand in demands:
                p = doc.add_paragraph(demand, style='List Number')
            
            doc.add_paragraph("")
            
            # Warning
            warning = doc.add_paragraph("TAKE NOTICE that if you fail to comply with the above demands within 15 days from the receipt of this notice, my client will be constrained to initiate appropriate legal proceedings against you for recovery of damages, compensation, and costs, without any further reference to you.")
            
            # Signature
            doc.add_paragraph("")
            doc.add_paragraph("Yours faithfully,")
            doc.add_paragraph("")
            doc.add_paragraph("[Advocate Name]")
            doc.add_paragraph("Advocate for the Notice Sender")
            doc.add_paragraph("[Address and Contact Details]")
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return {
                "success": True,
                "document": doc_bytes,
                "filename": f"Legal_Notice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "type": "legal_notice"
            }
            
        except Exception as e:
            logger.error(f"Legal notice generation error: {str(e)}")
            return {"success": False, "error": f"Template generation failed: {str(e)}"}
    
    def generate_court_petition(self, case_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a basic court petition template"""
        try:
            doc = Document()
            
            # Court header
            court_header = doc.add_heading('IN THE COURT OF [COURT NAME]', 0)
            court_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("")
            
            # Case details
            case_title = doc.add_heading(f'Case No. [To be filled by Court Registry]', 1)
            case_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("")
            
            # Parties
            doc.add_heading('BETWEEN:', 1)
            doc.add_paragraph("")
            
            doc.add_paragraph("[Petitioner Name]")
            doc.add_paragraph("[Petitioner Address]")
            doc.add_paragraph("...Petitioner")
            
            doc.add_paragraph("")
            doc.add_paragraph("AND")
            doc.add_paragraph("")
            
            doc.add_paragraph("[Respondent Name]")
            doc.add_paragraph("[Respondent Address]")
            doc.add_paragraph("...Respondent")
            
            doc.add_paragraph("")
            
            # Petition content
            doc.add_heading('PETITION', 1)
            doc.add_paragraph("")
            
            doc.add_paragraph("TO THE HONOURABLE COURT,")
            doc.add_paragraph("")
            
            doc.add_paragraph("The humble petition of the Petitioner above named MOST RESPECTFULLY SHOWETH:")
            doc.add_paragraph("")
            
            # Facts and grounds
            doc.add_heading('FACTS OF THE CASE:', 2)
            
            facts_para = doc.add_paragraph("1. ")
            facts_para.add_run(case_data.get('facts', '[Detailed facts of the case to be inserted here]'))
            
            doc.add_paragraph("")
            doc.add_paragraph("2. [Additional facts and chronology of events]")
            doc.add_paragraph("")
            doc.add_paragraph("3. [Legal grounds and cause of action]")
            doc.add_paragraph("")
            
            # Grounds for relief
            doc.add_heading('GROUNDS FOR RELIEF:', 2)
            doc.add_paragraph("a) [Legal ground 1]")
            doc.add_paragraph("b) [Legal ground 2]") 
            doc.add_paragraph("c) [Legal ground 3]")
            doc.add_paragraph("")
            
            # Prayer
            doc.add_heading('PRAYER:', 2)
            doc.add_paragraph("In the premises, it is most respectfully prayed that this Hon'ble Court may be pleased to:")
            doc.add_paragraph("")
            
            relief_type = case_data.get('relief_type', 'appropriate relief').replace('-', ' ').title()
            prayers = [
                f"Grant {relief_type} in favor of the Petitioner",
                "Award costs of this petition",
                "Grant such other relief as this Hon'ble Court may deem fit and proper"
            ]
            
            for i, prayer in enumerate(prayers, 1):
                doc.add_paragraph(f"{chr(ord('a') + i - 1)}) {prayer}")
            
            doc.add_paragraph("")
            
            # Verification
            doc.add_heading('VERIFICATION:', 2)
            doc.add_paragraph("I, [Petitioner Name], the Petitioner above named, do hereby verify that the contents of the above petition are true and correct to my knowledge and belief and that I have not concealed any material fact.")
            
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph("Place: [Place]")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("")
            doc.add_paragraph("                                                    [Petitioner Signature]")
            doc.add_paragraph("                                                    Petitioner")
            
            # Save document
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return {
                "success": True,
                "document": doc_bytes,
                "filename": f"Court_Petition_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "type": "court_petition"
            }
            
        except Exception as e:
            logger.error(f"Court petition generation error: {str(e)}")
            return {"success": False, "error": f"Petition generation failed: {str(e)}"}
    
    def generate_affidavit(self, case_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate an affidavit template"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('AFFIDAVIT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("")
            
            # Deponent details
            doc.add_paragraph("I, [Full Name], aged [Age] years, [occupation], residing at [Full Address], do hereby solemnly affirm and state as under:")
            doc.add_paragraph("")
            
            # Numbered statements
            doc.add_paragraph("1. That I am the [relationship to the case] in the above matter and am well acquainted with the facts and circumstances of the case.")
            doc.add_paragraph("")
            
            facts_para = doc.add_paragraph("2. That ")
            facts_para.add_run(case_data.get('facts', '[Facts and circumstances to be stated here]'))
            doc.add_paragraph("")
            
            doc.add_paragraph("3. That the facts stated above are true and correct to the best of my knowledge and belief.")
            doc.add_paragraph("")
            
            doc.add_paragraph("4. That I have not concealed any material fact and the same is being filed for the information of this Hon'ble Court.")
            doc.add_paragraph("")
            
            # Verification clause
            doc.add_paragraph("I hereby verify that the contents of this affidavit are true and correct to the best of my knowledge and belief and nothing material has been concealed therefrom.")
            doc.add_paragraph("")
            
            # Signature section
            doc.add_paragraph("")
            doc.add_paragraph("DEPONENT")
            doc.add_paragraph("")
            doc.add_paragraph("VERIFICATION:")
            doc.add_paragraph("Verified at [Place] on this [Date] day of [Month], [Year] that the contents of the above affidavit are true and correct to my knowledge and belief.")
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph("                                                    [Signature]")
            doc.add_paragraph("                                                    [Name of Deponent]")
            
            # Save document
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return {
                "success": True,
                "document": doc_bytes,
                "filename": f"Affidavit_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "type": "affidavit"
            }
            
        except Exception as e:
            logger.error(f"Affidavit generation error: {str(e)}")
            return {"success": False, "error": f"Affidavit generation failed: {str(e)}"}


def create_download_link(file_bytes: io.BytesIO, filename: str, link_text: str) -> str:
    """Create a download link for the generated document"""
    b64 = base64.b64encode(file_bytes.read()).decode()
    file_bytes.seek(0)  # Reset stream position
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" class="download-link">{link_text}</a>'
    return href


def display_document_upload_interface():
    """Display the document upload interface in Streamlit"""
    st.markdown('<div class="legal-details-header">📄 Document Processing & Analysis</div>', unsafe_allow_html=True)
    
    with st.container():
        st.markdown('<div class="legal-details-content">', unsafe_allow_html=True)
        
        # Document upload section
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📤 Upload Document")
            uploaded_file = st.file_uploader(
                "Choose a document",
                type=['pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg'],
                help="Supported formats: PDF, Word documents, Text files, Images"
            )
            
            if uploaded_file is not None:
                processor = DocumentProcessor()
                
                with st.spinner("🔍 Processing document..."):
                    result = processor.process_uploaded_file(uploaded_file)
                
                if result["success"]:
                    st.success("✅ Document processed successfully!")
                    
                    # Display extracted text
                    with st.expander("📋 Extracted Text", expanded=True):
                        st.text_area(
                            "Document Content",
                            value=result["text"][:2000] + ("..." if len(result["text"]) > 2000 else ""),
                            height=200,
                            help="First 2000 characters of extracted text"
                        )
                    
                    # Display metadata
                    with st.expander("ℹ️ Document Information"):
                        metadata = result["metadata"]
                        for key, value in metadata.items():
                            st.write(f"**{key.title()}:** {value}")
                    
                    # Option to use extracted text for analysis
                    if st.button("📊 Use This Text for Legal Analysis"):
                        st.session_state['document_text'] = result["text"]
                        st.success("✅ Text loaded for analysis! Scroll down to the case analysis form.")
                
                else:
                    st.error(f"❌ Error processing document: {result['error']}")
        
        with col2:
            st.markdown("#### 📝 Generate Legal Templates")
            
            template_type = st.selectbox(
                "Choose Template Type",
                options=['legal_notice', 'court_petition', 'affidavit'],
                format_func=lambda x: {
                    'legal_notice': '📄 Legal Notice',
                    'court_petition': '⚖️ Court Petition',
                    'affidavit': '📋 Affidavit'
                }[x]
            )
            
            if st.button("📄 Generate Template"):
                if 'relief' in st.session_state and 'facts' in st.session_state:
                    case_data = {
                        'relief_type': st.session_state.get('relief', 'legal-matter'),
                        'facts': st.session_state.get('facts', 'Case facts not provided')
                    }
                    
                    generator = LegalTemplateGenerator()
                    
                    with st.spinner("📝 Generating template..."):
                        if template_type == 'legal_notice':
                            result = generator.generate_legal_notice(case_data)
                        elif template_type == 'court_petition':
                            result = generator.generate_court_petition(case_data)
                        elif template_type == 'affidavit':
                            result = generator.generate_affidavit(case_data)
                    
                    if result["success"]:
                        st.success("✅ Template generated successfully!")
                        
                        # Create download link
                        download_link = create_download_link(
                            result["document"],
                            result["filename"],
                            f"📥 Download {template_type.replace('_', ' ').title()}"
                        )
                        st.markdown(download_link, unsafe_allow_html=True)
                        
                        st.info("💡 **Note:** This is a basic template. Please review and customize according to your specific case requirements and consult with a qualified lawyer.")
                    
                    else:
                        st.error(f"❌ Error generating template: {result['error']}")
                else:
                    st.warning("⚠️ Please complete the case analysis form first to generate personalized templates.")
        
        st.markdown('</div>', unsafe_allow_html=True)


# Additional CSS for download links
DOCUMENT_CSS = """
<style>
.download-link {
    display: inline-block;
    background: linear-gradient(135deg, var(--court-navy), var(--government-blue));
    color: var(--official-white) !important;
    padding: 0.8rem 1.5rem;
    border-radius: 6px;
    text-decoration: none;
    font-weight: 600;
    font-family: 'Inter', sans-serif;
    border: 2px solid var(--justice-gold);
    transition: all 0.2s ease;
    margin: 0.5rem 0;
}

.download-link:hover {
    background: linear-gradient(135deg, var(--government-blue), var(--court-navy));
    transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(30, 41, 81, 0.4);
    text-decoration: none;
    color: var(--official-white) !important;
}
</style>
"""
